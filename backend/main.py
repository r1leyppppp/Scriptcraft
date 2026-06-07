# =============================================================================
# Scriptcraft 后端服务
# 基于 FastAPI + 阿里云通义千问，提供小说转剧本及相关 AI 功能接口
# =============================================================================

# ===== 标准库 =====
import os
import io
import tempfile
import platform

# ===== 第三方库 =====
from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from openai import OpenAI
from docx import Document
from docx import Document as DocxDocument
import pypdf
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

# =============================================================================
# 应用初始化
# =============================================================================

app = FastAPI(title="Scriptcraft API", version="1.0.0")

# 允许跨域请求（前端直接打开 HTML 文件时需要）
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# =============================================================================
# 工具函数
# =============================================================================

def get_ai_client() -> OpenAI:
    """创建并返回阿里云通义千问 AI 客户端"""
    return OpenAI(
        api_key=os.environ.get("ALIBABA_API_KEY"),
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
    )

def ai_chat(prompt: str) -> str:
    """
    调用 AI 模型，返回文本响应
    
    Args:
        prompt: 用户输入的提示词
    
    Returns:
        AI 生成的文本内容
    """
    client = get_ai_client()
    response = client.chat.completions.create(
        model="qwen-plus",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

# =============================================================================
# 请求数据模型
# =============================================================================

class NovelInput(BaseModel):
    """小说转剧本请求体"""
    text: str

class EditInput(BaseModel):
    """剧本修改请求体"""
    script: str
    request: str

class ExportInput(BaseModel):
    """文件导出请求体"""
    content: str
    format: str
    filename: str

class ReviewInput(BaseModel):
    """剧本评审请求体"""
    script: str

class RelationshipInput(BaseModel):
    """人物关系分析请求体"""
    script: str

class StoryboardInput(BaseModel):
    """分镜脚本生成请求体"""
    script: str

class SvgExportInput(BaseModel):
    """SVG 导出请求体"""
    svg: str
    filename: str

# =============================================================================
# API 接口
# =============================================================================

# ----- 核心功能：小说转剧本 -----

@app.post("/convert")
async def convert(input: NovelInput):
    """
    将小说文本转换为 YAML 格式剧本
    调用 AI 以专业剧本改编师视角重新创作，输出符合 SCHEMA.md 定义的 YAML 结构
    """
    prompt = f"""你是一位专业的剧本改编师。请将以下小说文本改编为专业的剧本格式，用YAML结构输出。

改编要求：
1. 不要照抄原文，要用剧本语言重新创作
2. 场景描述要有画面感，像电影镜头一样
3. 人物动作要具体生动
4. 对话要自然流畅，符合人物性格
5. 添加必要的舞台/镜头指示

YAML格式要求：
- scene: 场景信息（地点、时间、氛围）
- characters: 人物列表（姓名、外貌、性格）
- actions: 动作描写
- dialogue: 对话（说话人、台词、动作）

小说原文：
{input.text}

请直接输出YAML内容，不要有任何解释文字。"""
    return {"result": ai_chat(prompt)}


@app.post("/edit")
async def edit(input: EditInput):
    """
    根据用户要求修改已有 YAML 剧本
    保持 YAML 格式不变，仅修改用户指定的内容
    """
    prompt = f"""你是一位专业的剧本改编师。以下是一份已生成的YAML格式剧本，用户希望对其进行修改。

当前剧本：
{input.script}

用户的修改要求：
{input.request}

请根据用户要求修改剧本，保持YAML格式输出，直接输出修改后的完整YAML内容，不要有任何解释文字。"""
    return {"result": ai_chat(prompt)}


# ----- 文件处理 -----

@app.post("/parse-file")
async def parse_file(file: UploadFile = File(...)):
    """
    解析上传的文件，提取纯文本内容
    支持格式：txt / docx / pdf
    """
    content = await file.read()
    filename = file.filename.lower()

    if filename.endswith('.txt'):
        text = content.decode('utf-8', errors='ignore')
    elif filename.endswith('.docx'):
        doc = Document(io.BytesIO(content))
        text = '\n'.join([para.text for para in doc.paragraphs])
    elif filename.endswith('.pdf'):
        reader = pypdf.PdfReader(io.BytesIO(content))
        text = '\n'.join([page.extract_text() for page in reader.pages])
    else:
        return {"error": "不支持的文件格式，请上传 txt / docx / pdf"}

    return {"text": text}


# ----- 导出功能 -----

@app.post("/export")
async def export_file(input: ExportInput):
    """
    将剧本内容导出为 DOCX 或 PDF 文件
    - DOCX：使用微软雅黑字体
    - PDF：注册中文字体，支持自动换行和分页
    """
    try:
        if input.format == 'docx':
            return _export_docx(input.content)
        elif input.format == 'pdf':
            return _export_pdf(input.content)
    except Exception as e:
        print(f"导出错误: {e}")
        return {"error": str(e)}


def _export_docx(content: str) -> StreamingResponse:
    """生成 DOCX 文件并返回流式响应"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    doc = DocxDocument()
    for line in content.split('\n'):
        para = doc.add_paragraph(line)
        run = para.runs[0] if para.runs else para.add_run()
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    return StreamingResponse(
        open(tmp.name, 'rb'),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': 'attachment; filename=script.docx'}
    )


def _export_pdf(content: str) -> StreamingResponse:
    """生成 PDF 文件并返回流式响应，支持中文字体和自动换行"""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 注册中文字体
    try:
        font_path = 'C:/Windows/Fonts/msyh.ttc' if platform.system() == 'Windows' \
            else '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc'
        pdfmetrics.registerFont(TTFont('Chinese', font_path))
        font_name = 'Chinese'
    except Exception:
        font_name = 'Helvetica'

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    c = canvas.Canvas(tmp.name, pagesize=A4)
    width, height = A4
    margin = 40
    max_width = width - margin * 2
    y = height - margin
    c.setFont(font_name, 12)

    def new_page():
        nonlocal y
        c.showPage()
        c.setFont(font_name, 12)
        y = height - margin

    # 逐行渲染，按字符宽度自动换行
    for line in content.split('\n'):
        if not line.strip():
            y -= 18
            if y < margin:
                new_page()
            continue
        current = ''
        for char in line:
            test = current + char
            if c.stringWidth(test, font_name, 12) > max_width:
                c.drawString(margin, y, current)
                y -= 18
                if y < margin:
                    new_page()
                current = char
            else:
                current = test
        if current:
            c.drawString(margin, y, current)
            y -= 18
            if y < margin:
                new_page()

    c.save()
    return StreamingResponse(
        open(tmp.name, 'rb'),
        media_type='application/pdf',
        headers={'Content-Disposition': 'attachment; filename=script.pdf'}
    )


# ----- AI 分析功能 -----

@app.post("/review")
async def review_script(input: ReviewInput):
    """
    对剧本进行专业评审
    返回 JSON 格式：总体评价、评分、优点、不足、建议
    """
    prompt = f"""你是一位专业的剧本评审导师。请对以下剧本进行专业评价，输出JSON格式。

剧本内容：
{input.script}

请输出以下JSON格式，不要有任何其他文字：
{{
  "overall": "总体评价描述",
  "score": 8,
  "strengths": ["优点1", "优点2", "优点3"],
  "weaknesses": ["不足1", "不足2"],
  "suggestions": ["建议1", "建议2", "建议3"]
}}"""
    return {"result": ai_chat(prompt)}


@app.post("/relationship")
async def get_relationship(input: RelationshipInput):
    """
    分析剧本中的人物关系
    返回 JSON 格式：人物列表 + 关系列表，用于前端生成关系图
    """
    prompt = f"""分析以下剧本中的人物关系，输出JSON格式。

剧本：
{input.script}

输出格式：
{{
  "characters": [
    {{"name": "人物名"}}
  ],
  "relations": [
    {{"from": "人物A", "to": "人物B", "type": "关系类型"}}
  ]
}}

只输出JSON，不要其他文字。"""
    return {"result": ai_chat(prompt)}


@app.post("/storyboard")
async def generate_storyboard(input: StoryboardInput):
    """
    根据剧本生成分镜脚本
    返回 YAML 格式：每个镜头包含景别、运动方式、画面描述、台词、时长
    """
    prompt = f"""你是专业的分镜师。根据以下剧本生成分镜脚本，用YAML结构输出。

YAML格式要求：
shots:
  - id: 1
    scene: 场景名
    shot_type: 全景/中景/近景/特写
    movement: 固定/推镜/拉镜/摇镜/跟镜
    description: 画面描述
    dialogue: 台词（无则留空）
    duration: 3

只输出YAML内容，不要任何解释文字。

剧本内容：
{input.script}"""
    return {"result": ai_chat(prompt)}


@app.post("/export-svg")
async def export_svg(input: SvgExportInput):
    """
    将 SVG 人物关系图导出为 PDF（备用接口，前端目前使用 Canvas 导出 PNG）
    """
    try:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        c = canvas.Canvas(tmp.name, pagesize=A4)
        width, height = A4
        c.setFont("Helvetica", 10)
        y = height - 40
        for line in input.svg.split('\n'):
            if y < 40:
                c.showPage()
                y = height - 40
            c.drawString(40, y, line[:100])
            y -= 14
        c.save()
        return StreamingResponse(
            open(tmp.name, 'rb'),
            media_type='application/pdf',
            headers={'Content-Disposition': 'attachment; filename=relationship.pdf'}
        )
    except Exception as e:
        print(f"SVG导出错误: {e}")
        return {"error": str(e)}